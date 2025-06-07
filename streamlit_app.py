import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import csv
from docx import Document
from ec2_sql_sizing import EC2DatabaseSizingCalculator

class EC2SizingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AWS EC2 Sizing Calculator")
        self.root.geometry("800x600")
        self.calculator = EC2DatabaseSizingCalculator()
        self.results = {}
        self.create_widgets()

    def create_widgets(self):
        # Frame for inputs
        input_frame = ttk.LabelFrame(self.root, text="Input Parameters")
        input_frame.pack(fill="x", padx=10, pady=10)

        # Input fields
        self.entries = {}
        fields = [
            ("on_prem_cores", "On-Prem CPU Cores", 16),
            ("peak_cpu_percent", "Peak CPU Utilization (%)", 65),
            ("on_prem_ram_gb", "On-Prem RAM (GB)", 64),
            ("peak_ram_percent", "Peak RAM Utilization (%)", 75),
            ("storage_current_gb", "Current DB Storage (GB)", 500),
            ("storage_growth_rate", "Annual Storage Growth Rate (e.g., 0.15)", 0.15),
            ("peak_iops", "Peak IOPS", 8000),
            ("peak_throughput_mbps", "Peak Throughput (MB/s)", 400),
            ("years", "Growth Projection (Years)", 3)
        ]

        for i, (key, label, default) in enumerate(fields):
            ttk.Label(input_frame, text=label).grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ttk.Entry(input_frame)
            entry.insert(0, str(default))
            entry.grid(row=i, column=1, padx=5, pady=2, sticky="ew")
            self.entries[key] = entry

        # Workload Profile Dropdown
        ttk.Label(input_frame, text="Workload Profile").grid(row=len(fields), column=0, sticky="w", padx=5, pady=2)
        self.workload_var = tk.StringVar()
        self.workload_combo = ttk.Combobox(input_frame, textvariable=self.workload_var, state="readonly")
        self.workload_combo['values'] = ("general", "memory", "compute")
        self.workload_combo.current(0)
        self.workload_combo.grid(row=len(fields), column=1, padx=5, pady=2, sticky="ew")

        # AMD Preference Checkbox
        self.amd_var = tk.BooleanVar(value=True)
        amd_check = ttk.Checkbutton(input_frame, text="Prefer AMD Instances (Cost Optimized)", variable=self.amd_var)
        amd_check.grid(row=len(fields)+1, column=0, columnspan=2, pady=5)

        # Buttons
        button_frame = ttk.Frame(self.root)
        button_frame.pack(pady=10)

        ttk.Button(button_frame, text="Generate Recommendations", command=self.calculate).grid(row=0, column=0, padx=10)
        ttk.Button(button_frame, text="Export CSV", command=self.export_csv).grid(row=0, column=1, padx=10)
        ttk.Button(button_frame, text="Export DOCX", command=self.export_docx).grid(row=0, column=2, padx=10)

        # Output Display
        self.tree = ttk.Treeview(self.root, columns=("Instance", "vCPUs", "RAM", "Storage", "EBS", "IOPS", "Throughput", "Processor"), show="headings")
        for col in self.tree["columns"]:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=100, anchor="center")
        self.tree.pack(fill="both", expand=True, padx=10, pady=10)

    def calculate(self):
        try:
            for key, entry in self.entries.items():
                val = float(entry.get()) if "." in entry.get() else int(entry.get())
                self.calculator.inputs[key] = val

            self.calculator.inputs["workload_profile"] = self.workload_var.get()
            self.calculator.inputs["prefer_amd"] = self.amd_var.get()
            self.results = self.calculator.generate_all_recommendations()

            self.tree.delete(*self.tree.get_children())
            for env, rec in self.results.items():
                self.tree.insert("", "end", values=(
                    f"{env}: {rec['instance_type']}", rec["vCPUs"], rec["RAM_GB"],
                    rec["storage_GB"], rec["ebs_type"], rec["iops_required"],
                    rec["throughput_required"], rec["processor"]
                ))
            messagebox.showinfo("Success", "Recommendations generated successfully.")

        except Exception as e:
            messagebox.showerror("Input Error", str(e))

    def export_csv(self):
        if not self.results:
            messagebox.showwarning("No Data", "Please generate results first.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])
        if path:
            with open(path, "w", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(["Environment", "Instance Type", "vCPUs", "RAM_GB", "Storage_GB", "EBS Type", "IOPS", "Throughput", "Processor"])
                for env, rec in self.results.items():
                    writer.writerow([env, rec["instance_type"], rec["vCPUs"], rec["RAM_GB"], rec["storage_GB"], rec["ebs_type"], rec["iops_required"], rec["throughput_required"], rec["processor"]])
            messagebox.showinfo("Exported", f"Results exported to {path}")

    def export_docx(self):
        if not self.results:
            messagebox.showwarning("No Data", "Please generate results first.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if path:
            doc = Document()
            doc.add_heading('EC2 Sizing Recommendations', 0)
            for env, rec in self.results.items():
                doc.add_heading(f'{env} Environment', level=1)
                for key, val in rec.items():
                    doc.add_paragraph(f"{key}: {val}")
            doc.save(path)
            messagebox.showinfo("Exported", f"Results exported to {path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = EC2SizingApp(root)
    root.mainloop()
